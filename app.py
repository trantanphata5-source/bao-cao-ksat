import os
import json
import uuid
import re
from datetime import datetime
from flask import Flask, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from docx import Document

app = Flask(__name__, static_folder='static', static_url_path='')
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['REPORTS_FOLDER'] = 'reports'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB max

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['REPORTS_FOLDER'], exist_ok=True)

def parse_docx(file_path, filename):
    doc = Document(file_path)
    
    report_data = {
        'id': str(uuid.uuid4()),
        'filename': filename,
        'upload_date': datetime.now().isoformat(),
        'date_range': '',
        'section1': [], # Kết quả triển khai
        'section2': [], # Kế hoạch
        'section3': [], # Kiến nghị
        'table_data': []
    }
    
    current_section = 0 # 0: header, 1: Muc I, 2: Muc II, 3: Muc III
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        # Try to extract date range from early paragraphs
        if current_section == 0 and '(từ ngày' in text.lower():
            report_data['date_range'] = text
            
        # Detect sections
        lower_text = text.lower()
        if re.match(r'^i\.\s+báo cáo kết quả', lower_text) or 'báo cáo kết quả triển khai công tác' in lower_text:
            current_section = 1
            report_data['section1'].append({'text': text, 'bold': True})
            continue
        elif re.match(r'^ii\.\s+kế hoạch', lower_text) or 'kế hoạch công tác từ ngày' in lower_text:
            current_section = 2
            report_data['section2'].append({'text': text, 'bold': True})
            continue
        elif re.match(r'^iii\.\s+kiến nghị', lower_text) or 'kiến nghị về công tác' in lower_text:
            current_section = 3
            report_data['section3'].append({'text': text, 'bold': True})
            continue
            
        is_bold = any(r.bold for r in p.runs if r.bold)
        
        item = {'text': text, 'bold': is_bold}
        
        if current_section == 1:
            report_data['section1'].append(item)
        elif current_section == 2:
            report_data['section2'].append(item)
        elif current_section == 3:
            report_data['section3'].append(item)

    # Extract tables
    for table in doc.tables:
        # Check if it's the data table (has headers like STT, Đơn vị, Kế hoạch...)
        if len(table.rows) > 1 and len(table.columns) >= 7:
            headers = [c.text.strip().replace('\n', ' ') for c in table.rows[0].cells]
            if 'STT' in headers[0] and 'Đơn vị' in headers[1]:
                for ri, row in enumerate(table.rows):
                    if ri == 0: continue # Skip header
                    row_data = [c.text.strip().replace('\n', ' ') for c in row.cells]
                    if len(row_data) >= 7 and row_data[0].isdigit():
                        try:
                            # 0: STT, 1: Don vi, 2: Ke hoach, 3: TB, 4: TH, 5: Luy ke, 6: Con lai
                            report_data['table_data'].append({
                                'stt': row_data[0],
                                'don_vi': row_data[1],
                                'ke_hoach': int(row_data[2]) if row_data[2].isdigit() else 0,
                                'trung_binh': int(row_data[3]) if row_data[3].isdigit() else 0,
                                'thuc_hien': int(row_data[4]) if row_data[4].isdigit() else 0,
                                'luy_ke': int(row_data[5]) if row_data[5].isdigit() else 0,
                                'con_lai': int(row_data[6]) if row_data[6].isdigit() else 0
                            })
                        except ValueError:
                            pass
                break # Found the main table, ignore others (like header/footer tables)
    
    # Generate a title
    report_data['title'] = report_data['date_range'] if report_data['date_range'] else f"Báo cáo ngày {datetime.now().strftime('%d/%m/%Y')}"
    
    return report_data


@app.route('/')
def serve_index():
    return send_from_directory('static', 'index.html')

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
        
    if file and file.filename.endswith('.docx'):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        try:
            report_data = parse_docx(filepath, filename)
            
            # Save to JSON
            json_path = os.path.join(app.config['REPORTS_FOLDER'], f"{report_data['id']}.json")
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
                
            return jsonify({'success': True, 'report': report_data})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
            
    return jsonify({'error': 'Invalid file format. Only .docx is supported.'}), 400


@app.route('/api/reports', methods=['GET'])
def get_reports():
    reports = []
    for filename in os.listdir(app.config['REPORTS_FOLDER']):
        if filename.endswith('.json'):
            with open(os.path.join(app.config['REPORTS_FOLDER'], filename), 'r', encoding='utf-8') as f:
                data = json.load(f)
                reports.append({
                    'id': data['id'],
                    'title': data.get('title', ''),
                    'date_range': data.get('date_range', ''),
                    'upload_date': data.get('upload_date', '')
                })
    # Sort by upload date descending
    reports.sort(key=lambda x: x['upload_date'], reverse=True)
    return jsonify(reports)

@app.route('/api/report/<report_id>', methods=['GET'])
def get_report(report_id):
    filepath = os.path.join(app.config['REPORTS_FOLDER'], f"{report_id}.json")
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            return jsonify(json.load(f))
    return jsonify({'error': 'Report not found'}), 404

@app.route('/api/report/<report_id>', methods=['DELETE'])
def delete_report(report_id):
    filepath = os.path.join(app.config['REPORTS_FOLDER'], f"{report_id}.json")
    if os.path.exists(filepath):
        os.remove(filepath)
        return jsonify({'success': True})
    return jsonify({'error': 'Report not found'}), 404

if __name__ == '__main__':
    app.run(debug=True, port=5000)
